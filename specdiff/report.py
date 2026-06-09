"""Report assembly and export generation."""

import logging
from pathlib import Path

import fitz  # PyMuPDF

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

from specdiff.models import (
    Change,
    ChangeKind,
    ComparisonResult,
    ComparisonSummary,
    DocumentMetadata,
)

logger = logging.getLogger(__name__)


def assemble_result(
    old_metadata: DocumentMetadata,
    new_metadata: DocumentMetadata,
    text_changes: list[Change],
    table_changes: list[Change],
    figure_changes: list[Change],
    manual_review_pages: list[int],
) -> ComparisonResult:
    """Assemble complete comparison result with summary."""
    all_changes = text_changes + table_changes + figure_changes

    # Compute summary
    summary = ComparisonSummary(
        insertions=sum(1 for c in all_changes if c.kind == ChangeKind.INSERTION),
        deletions=sum(1 for c in all_changes if c.kind == ChangeKind.DELETION),
        modifications=sum(1 for c in all_changes if c.kind == ChangeKind.MODIFICATION),
        table_changes=sum(1 for c in all_changes if c.kind == ChangeKind.TABLE_CHANGE),
        figure_flags=sum(1 for c in all_changes if c.kind == ChangeKind.FIGURE_FLAG),
        pages_changed=sorted(set(c.location.page for c in all_changes)),
    )

    # Collect warnings
    warnings: list[str] = []

    if figure_changes:
        warnings.append(
            f"{len(figure_changes)} pages contain figures/images that require manual review"
        )

    if old_metadata.page_count != new_metadata.page_count:
        warnings.append(
            f"Page count changed from {old_metadata.page_count} to {new_metadata.page_count}"
        )

    logger.info(
        f"Assembled result: {len(all_changes)} total changes, "
        f"{summary.pages_changed.__len__()} pages affected"
    )

    return ComparisonResult(
        old_metadata=old_metadata,
        new_metadata=new_metadata,
        changes=all_changes,
        summary=summary,
        manual_review_pages=manual_review_pages,
        warnings=warnings,
    )


def export_docx(result: ComparisonResult, output_path: Path) -> None:
    """Export comparison result as Word document with tracked changes style."""
    doc = Document()

    # Title
    doc.add_heading("Specification Comparison Report", 0)

    # Summary
    doc.add_heading("Summary", 1)
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = "Light Grid Accent 1"

    summary_data = [
        ("Pages Changed", str(len(result.summary.pages_changed))),
        ("Insertions", str(result.summary.insertions)),
        ("Deletions", str(result.summary.deletions)),
        ("Modifications", str(result.summary.modifications)),
        ("Table Changes", str(result.summary.table_changes)),
        ("Figure Flags", str(result.summary.figure_flags)),
    ]

    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    # Warnings
    if result.warnings:
        doc.add_heading("Warnings", 2)
        for warning in result.warnings:
            p = doc.add_paragraph(warning, style="List Bullet")
            # Highlight warnings
            for run in p.runs:
                run.font.color.rgb = RGBColor(184, 134, 11)

    # Manual review
    if result.manual_review_pages:
        doc.add_heading("⚠️ Manual Review Required", 2)
        doc.add_paragraph(
            "The following pages contain figures or images that may have changed. "
            "Please review manually:"
        )
        for page in result.manual_review_pages:
            doc.add_paragraph(f"Page {page + 1}", style="List Bullet")

    # Changes
    doc.add_heading("Detected Changes", 1)

    for change in result.changes:
        # Location header
        loc_parts = []
        if change.location.clause_id:
            loc_parts.append(f"Clause {change.location.clause_id}")
        loc_parts.append(f"Page {change.location.page + 1}")
        if change.location.table_ref:
            loc_parts.append(change.location.table_ref)
        if change.location.cell_ref:
            loc_parts.append(f"Cell {change.location.cell_ref}")

        p = doc.add_paragraph()
        loc_run = p.add_run(" | ".join(loc_parts))
        loc_run.font.color.rgb = RGBColor(128, 128, 128)
        loc_run.font.size = 10

        # Change content
        p = doc.add_paragraph()

        if change.kind == ChangeKind.DELETION:
            run = p.add_run(change.old_text)
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)

        elif change.kind == ChangeKind.INSERTION:
            run = p.add_run(change.new_text)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.color.rgb = RGBColor(0, 128, 0)

        else:
            if change.old_text:
                run = p.add_run(change.old_text)
                run.font.strike = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                p.add_run(" → ")

            if change.new_text:
                run = p.add_run(change.new_text)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.color.rgb = RGBColor(0, 128, 0)

    doc.save(output_path)
    logger.info(f"Exported DOCX to {output_path}")


def export_pdf(result: ComparisonResult, output_path: Path) -> None:
    """Export comparison result as a PDF report using PyMuPDF."""
    doc = fitz.open()

    # Colours
    RED = (0.8, 0.0, 0.0)
    GREEN = (0.0, 0.55, 0.0)
    BLUE = (0.0, 0.33, 0.67)
    ORANGE = (1.0, 0.55, 0.0)
    PURPLE = (0.55, 0.0, 0.6)
    GREY = (0.4, 0.4, 0.4)
    BLACK = (0.0, 0.0, 0.0)
    LIGHT_YELLOW = (1.0, 0.98, 0.82)
    WHITE = (1.0, 1.0, 1.0)

    KIND_COLOUR = {
        ChangeKind.INSERTION: GREEN,
        ChangeKind.DELETION: RED,
        ChangeKind.MODIFICATION: BLUE,
        ChangeKind.TABLE_CHANGE: ORANGE,
        ChangeKind.FIGURE_FLAG: PURPLE,
    }
    KIND_LABEL = {
        ChangeKind.INSERTION: "INSERTION",
        ChangeKind.DELETION: "DELETION",
        ChangeKind.MODIFICATION: "MODIFICATION",
        ChangeKind.TABLE_CHANGE: "TABLE CHANGE",
        ChangeKind.FIGURE_FLAG: "FIGURE FLAG",
    }

    PAGE_W, PAGE_H = 595, 842  # A4 points
    MARGIN = 50
    BODY_W = PAGE_W - 2 * MARGIN
    FONT = "helv"
    FONT_BOLD = "hebo"
    LINE_GAP = 4  # extra space between lines

    def new_page() -> tuple[fitz.Page, float]:
        page = doc.new_page(width=PAGE_W, height=PAGE_H)
        return page, float(MARGIN)

    def check_y(page: fitz.Page, y: float, needed: float = 20) -> tuple[fitz.Page, float]:
        """Return a new page if there isn't enough vertical space remaining."""
        if y + needed > PAGE_H - MARGIN:
            page, y = new_page()
        return page, y

    # ── y is always the TOP of the current text line ─────────────────────────
    # PyMuPDF insert_text takes the baseline, so baseline = y_top + fontsize
    # After drawing, advance by fontsize + LINE_GAP

    def write_line(
        page: fitz.Page,
        y: float,           # top of the line
        text: str,
        size: float = 10,
        colour: tuple = BLACK,
        bold: bool = False,
        indent: float = 0,
    ) -> float:
        """Draw one line of text. y = top of line. Returns y of next line top."""
        font = FONT_BOLD if bold else FONT
        baseline = y + size          # baseline is below the top by 1 em
        page.insert_text(
            (MARGIN + indent, baseline),
            text,
            fontname=font,
            fontsize=size,
            color=colour,
        )
        return y + size + LINE_GAP   # advance past glyph + gap

    def write_wrapped(
        page: fitz.Page,
        y: float,
        text: str,
        size: float = 9,
        colour: tuple = BLACK,
        indent: float = 0,
        max_chars: int = 88,
    ) -> tuple[fitz.Page, float]:
        """Word-wrap text across lines/pages. y = top of first line. Returns (page, y_next)."""
        line_h = size + LINE_GAP
        words = text.split()
        current = ""
        for word in words:
            candidate = (current + " " + word).strip()
            if len(candidate) > max_chars and current:
                page, y = check_y(page, y, line_h)
                y = write_line(page, y, current, size=size, colour=colour, indent=indent)
                current = word
            else:
                current = candidate
        if current:
            page, y = check_y(page, y, line_h)
            y = write_line(page, y, current, size=size, colour=colour, indent=indent)
        return page, y

    # ── Cover / Summary page ─────────────────────────────────────────────────
    page, y = new_page()
    y += 10

    y = write_line(page, y, "Specification Comparison Report", size=18, bold=True)
    y += 4
    y = write_line(page, y, f"Old:  {result.old_metadata.file_hash[:12]}...", size=9, colour=GREY)
    y = write_line(page, y, f"New:  {result.new_metadata.file_hash[:12]}...", size=9, colour=GREY)
    y += 10

    # Build stats list first so we can size the box
    stats = [
        ("Pages changed",  str(len(result.summary.pages_changed))),
        ("Insertions",     str(result.summary.insertions)),
        ("Deletions",      str(result.summary.deletions)),
        ("Modifications",  str(result.summary.modifications)),
        ("Table changes",  str(result.summary.table_changes)),
        ("Figure flags",   str(result.summary.figure_flags)),
    ]
    ROW_H = 10 + LINE_GAP           # fontsize + gap
    box_h = ROW_H + 6 + len(stats) * ROW_H + 10   # title + padding + rows + bottom pad
    page.draw_rect(
        fitz.Rect(MARGIN, y, MARGIN + BODY_W, y + box_h),
        color=(0.85, 0.85, 0.85),
        fill=LIGHT_YELLOW,
    )
    y += 6
    y = write_line(page, y, "Summary", size=12, bold=True, indent=6)
    for label, val in stats:
        page, y = check_y(page, y, ROW_H + 2)
        y = write_line(page, y, f"{label}:  {val}", size=10, indent=12)
    y += 8

    # Warnings
    if result.warnings:
        y += 4
        page, y = check_y(page, y, 20)
        y = write_line(page, y, "Warnings", size=11, bold=True, colour=ORANGE)
        for w in result.warnings:
            page, y = check_y(page, y, 14)
            page, y = write_wrapped(page, y, f"  {w}", size=9, colour=ORANGE, indent=6)
        y += 4

    # Manual review
    if result.manual_review_pages:
        y += 6
        page, y = check_y(page, y, 24)
        y = write_line(page, y, "Manual Review Required", size=11, bold=True, colour=RED)
        pages_str = ", ".join(str(p + 1) for p in result.manual_review_pages)
        page, y = write_wrapped(page, y, f"Pages: {pages_str}", size=9, colour=RED, indent=6)
        y += 4

    # ── Changes section ───────────────────────────────────────────────────────
    y += 10
    page, y = check_y(page, y, 30)
    y = write_line(page, y, f"Detected Changes  ({len(result.changes)} total)", size=13, bold=True)
    y += 6

    BAR_H = 18   # coloured badge bar height — tall enough for 8pt text + padding

    for change in result.changes:
        colour = KIND_COLOUR.get(change.kind, BLACK)
        label  = KIND_LABEL.get(change.kind, change.kind.value)

        # Need at least bar + one text line before forcing a new page
        page, y = check_y(page, y, BAR_H + 30)

        # ── Coloured badge bar ────────────────────────────────────────────────
        bar_top = y
        page.draw_rect(
            fitz.Rect(MARGIN, bar_top, MARGIN + BODY_W, bar_top + BAR_H),
            color=colour,
            fill=colour,
        )
        # Text inside the bar: baseline = bar_top + 8pt + 3pt top-padding
        bar_baseline = bar_top + 3 + 9   # 3 px padding + 9pt font
        page.insert_text(
            (MARGIN + 6, bar_baseline),
            label,
            fontname=FONT_BOLD,
            fontsize=9,
            color=WHITE,
        )
        # Location info on right side of bar
        loc_parts = []
        if change.location.clause_id:
            loc_parts.append(f"Clause {change.location.clause_id}")
        loc_parts.append(f"Page {change.location.page + 1}")
        if change.location.table_ref:
            loc_parts.append(change.location.table_ref)
        if change.location.cell_ref:
            loc_parts.append(f"Cell {change.location.cell_ref}")
        page.insert_text(
            (MARGIN + 110, bar_baseline),
            "  |  ".join(loc_parts),
            fontname=FONT,
            fontsize=9,
            color=WHITE,
        )

        # y now starts BELOW the bar — clean separation
        y = bar_top + BAR_H + 6

        # ── Old text ─────────────────────────────────────────────────────────
        if change.old_text:
            page, y = check_y(page, y, 14)
            y = write_line(page, y, "REMOVED:", size=8, colour=RED, bold=True, indent=6)
            page, y = write_wrapped(page, y, change.old_text, size=9, colour=RED, indent=16)
            y += 3

        # ── New text ─────────────────────────────────────────────────────────
        if change.new_text:
            page, y = check_y(page, y, 14)
            y = write_line(page, y, "ADDED:", size=8, colour=GREEN, bold=True, indent=6)
            page, y = write_wrapped(page, y, change.new_text, size=9, colour=GREEN, indent=16)
            y += 3

        y += 8   # gap between change blocks

    doc.save(str(output_path))
    doc.close()
    logger.info(f"Exported PDF to {output_path}")


def export_all(result: ComparisonResult, output_dir: Path, base_name: str = "comparison") -> None:
    """Export result as PDF and DOCX only."""
    output_dir.mkdir(parents=True, exist_ok=True)

    export_docx(result, output_dir / f"{base_name}.docx")
    export_pdf(result, output_dir / f"{base_name}.pdf")

    logger.info(f"All exports complete in {output_dir}")
